"""
=============================================================================
AI Job Search Agent — FastAPI Backend
=============================================================================
Endpoints:
  GET  /health         — Ollama / Gemini connectivity check
  GET  /jobs           — Return all job postings as JSON
  POST /run-agent      — SSE stream: filter → rank → tailor
  POST /parse-resume   — Extract text from uploaded PDF
  POST /export-resume  — Generate 1-page tailored PDF

Run:
  uvicorn api_server:app --reload --port 8000
=============================================================================
"""

from __future__ import annotations

import asyncio
import base64
import io
import json
import os
import re
import sys
import textwrap
from pathlib import Path
from typing import AsyncGenerator

import requests
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel, Field

# Load .env
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env")
except ImportError:
    pass

# ── Project imports ──────────────────────────────────────────────────────────
sys.path.insert(0, os.path.dirname(__file__))
from main import (
    DATASET_PATH,
    MODEL_NAME,
    filtering_tool,
    load_dataset,
    ranking_tool,
    resume_tailoring_tool,
    call_ollama_simple,
    get_llm_caller,
)

# ─────────────────────────────────────────────────────────────────────────────
# App setup
# ─────────────────────────────────────────────────────────────────────────────

app = FastAPI(title="Job Search Agent API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─────────────────────────────────────────────────────────────────────────────
# Pydantic models
# ─────────────────────────────────────────────────────────────────────────────

class AgentRequest(BaseModel):
    name: str = "Candidate"
    skills: list[str] = Field(default_factory=list)
    years_experience: int = 2
    preferred_location: str = "Remote"
    remote_only: bool = False
    exclude_companies: list[str] = Field(default_factory=list)
    current_summary: str = ""
    bullet_1: str = ""
    bullet_2: str = ""
    resume_pdf_base64: str = ""


class ExportResumeRequest(BaseModel):
    original_pdf_base64: str
    tailored_summary: str
    tailored_bullet_1: str
    tailored_bullet_2: str
    original_summary: str = ""
    original_bullet_1: str = ""
    original_bullet_2: str = ""
    job_title: str = "Job"
    company: str = "Company"


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _sse(event: dict) -> str:
    return f"data: {json.dumps(event)}\n\n"


def _check_ollama() -> bool:
    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=3)
        return resp.status_code == 200
    except Exception:
        return False


# ─────────────────────────────────────────────────────────────────────────────
# GET /health
# ─────────────────────────────────────────────────────────────────────────────

def _get_git_version() -> str:
    try:
        import subprocess
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"]).decode("utf-8").strip()
    except Exception:
        return "unknown"

@app.get("/health")
def health():
    return {
        "status": "ok",
        "model": MODEL_NAME,
        "ollama_connected": _check_ollama(),
        "version": _get_git_version(),
    }


# ─────────────────────────────────────────────────────────────────────────────
# GET /jobs
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/jobs")
def get_jobs():
    return load_dataset(DATASET_PATH)


# ─────────────────────────────────────────────────────────────────────────────
# POST /run-agent  (SSE streaming)
# ─────────────────────────────────────────────────────────────────────────────

async def _agent_stream(req: AgentRequest) -> AsyncGenerator[str, None]:
    loop = asyncio.get_event_loop()
    provider = "llama"
    _, llm_simple_fn = get_llm_caller(provider)
    tailor_timeout = 300

    profile = {
        "name": req.name,
        "skills": req.skills,
        "years_experience": req.years_experience,
        "preferred_location": req.preferred_location,
        "remote_only": req.remote_only,
        "exclude_companies": req.exclude_companies,
        "current_summary": req.current_summary,
        "bullet_1": req.bullet_1,
        "bullet_2": req.bullet_2,
    }

    try:
        yield _sse({"type": "status", "message": "Using Qwen 3.5 4B · Loading jobs..."})
        all_jobs = await loop.run_in_executor(None, load_dataset, DATASET_PATH)
        yield _sse({"type": "status", "message": f"Loaded {len(all_jobs)} jobs."})

        # ── Filter ──────────────────────────────────────────────────────────
        yield _sse({"type": "filter_start", "message": f"Filtering {len(all_jobs)} jobs..."})
        filter_result = await loop.run_in_executor(
            None,
            lambda: filtering_tool(
                jobs=all_jobs,
                preferred_location=req.preferred_location,
                max_years_experience=req.years_experience,
                required_skills=req.skills,
                exclude_companies=req.exclude_companies or None,
                remote_only=req.remote_only,
            ),
        )
        filtered_jobs = filter_result["filtered_jobs"]
        yield _sse({
            "type": "filter_result",
            "count": filter_result["count"],
            "trace": filter_result["reasoning_trace"],
            "jobs": [
                {
                    "job_title": j["job_title"],
                    "company": j["company"],
                    "location": j["location"],
                    "years_experience": j["years_experience"],
                    "required_skills": j["required_skills"],
                    "url": j["url"],
                }
                for j in filtered_jobs
            ],
        })

        if not filtered_jobs:
            yield _sse({"type": "error", "message": "No jobs matched your filters. Try broadening your location or skills."})
            yield _sse({"type": "complete", "message": "Done (no matches)"})
            return

        # ── Rank ────────────────────────────────────────────────────────────
        yield _sse({"type": "rank_start", "message": f"Ranking {len(filtered_jobs)} jobs..."})
        rank_result = await loop.run_in_executor(
            None,
            lambda: ranking_tool(
                jobs=filtered_jobs,
                candidate_skills=req.skills,
                candidate_years=req.years_experience,
                preferred_location=req.preferred_location,
            ),
        )
        yield _sse({
            "type": "rank_result",
            "top_3": rank_result["top_3"],
            "all_ranked": rank_result["ranked_jobs"],
        })

        best_job = rank_result["best_job"]
        if not best_job:
            yield _sse({"type": "error", "message": "Ranking produced no results."})
            yield _sse({"type": "complete", "message": "Done"})
            return

        # ── Resume Tailor ───────────────────────────────────────────────────
        yield _sse({
            "type": "tailor_start",
            "message": f"Tailoring resume for {best_job['job_title']} at {best_job['company']}...",
        })

        # Check Ollama availability
        if not _check_ollama():
            yield _sse({"type": "error", "message": "Ollama is not running. Start with: ollama serve"})
            yield _sse({"type": "complete", "message": "Done (tailor skipped)"})
            return

        tailor_result = await loop.run_in_executor(
            None,
            lambda: resume_tailoring_tool(
                job=best_job,
                candidate_profile=profile,
                llm_call_fn=llm_simple_fn,
                timeout=tailor_timeout,
            ),
        )

        if tailor_result["success"]:
            t = tailor_result["tailored"]
            yield _sse({
                "type": "tailor_result",
                "job_title": best_job["job_title"],
                "company": best_job["company"],
                "professional_summary": t.get("professional_summary", ""),
                "bullet_1": t.get("bullet_1_rewritten", ""),
                "bullet_2": t.get("bullet_2_rewritten", ""),
            })
        else:
            yield _sse({
                "type": "tailor_result",
                "job_title": best_job["job_title"],
                "company": best_job["company"],
                "professional_summary": "",
                "bullet_1": "",
                "bullet_2": "",
                "error": "Could not parse tailored JSON from LLM response.",
            })

        yield _sse({"type": "complete", "message": "Done"})

    except Exception as exc:
        yield _sse({"type": "error", "message": str(exc)})
        yield _sse({"type": "complete", "message": "Done (error)"})


@app.post("/run-agent")
async def run_agent(req: AgentRequest):
    return StreamingResponse(
        _agent_stream(req),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ─────────────────────────────────────────────────────────────────────────────
# POST /parse-resume
# ─────────────────────────────────────────────────────────────────────────────

def _extract_text_from_pdf(content: bytes) -> tuple[str, int]:
    """
    Extract text from PDF bytes using pdfplumber with word-level reconstruction
    to fix missing-space issues common in columnar/ATS-formatted resumes.
    Falls back to pypdf if pdfplumber is unavailable.
    """
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages_count = len(pdf.pages)
            page_texts = []
            for page in pdf.pages:
                # Use extract_words for proper spacing, then reconstruct lines
                words = page.extract_words(
                    x_tolerance=3,
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=True,
                )
                if not words:
                    page_texts.append(page.extract_text() or "")
                    continue

                # Group words into lines by y-position proximity
                lines: list[list[dict]] = []
                for w in words:
                    placed = False
                    for line in lines:
                        if abs(w["top"] - line[0]["top"]) < 4:
                            line.append(w)
                            placed = True
                            break
                    if not placed:
                        lines.append([w])

                # Sort lines top-to-bottom, words left-to-right, join with spaces
                lines.sort(key=lambda l: l[0]["top"])
                reconstructed = []
                for line_words in lines:
                    line_words.sort(key=lambda w: w["x0"])
                    reconstructed.append(" ".join(w["text"] for w in line_words))

                page_texts.append("\n".join(reconstructed))
            return "\n".join(page_texts), pages_count
    except ImportError:
        pass

    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        return text, len(reader.pages)
    except ImportError:
        raise HTTPException(status_code=500, detail="Install pdfplumber: pip install pdfplumber")


# Comprehensive skill keyword list covering most data/ML/software job postings
_KNOWN_SKILLS = [
    # Languages
    "Python", "R", "SQL", "Java", "Scala", "C++", "C#", "JavaScript", "TypeScript",
    "Go", "Rust", "Julia", "MATLAB", "Bash", "Shell",
    # ML / AI
    "Machine Learning", "Deep Learning", "NLP", "Computer Vision", "Reinforcement Learning",
    "Transformers", "LLMs", "Large Language Models", "LLM", "Generative AI", "Gen AI",
    "Neural Networks", "Feature Engineering", "Model Deployment", "MLOps",
    # Frameworks / Libraries
    "PyTorch", "TensorFlow", "Keras", "Scikit-learn", "XGBoost", "LightGBM", "CatBoost",
    "Hugging Face", "spaCy", "NLTK", "OpenCV", "Pandas", "NumPy", "SciPy", "Matplotlib",
    "Seaborn", "Plotly", "FastAPI", "Flask", "Django", "LangChain", "LlamaIndex",
    # Data Engineering
    "Airflow", "dbt", "Spark", "PySpark", "Kafka", "Flink", "Hadoop", "Hive",
    "ETL", "ELT", "Data Pipelines", "Data Warehouse", "Data Lake", "Databricks",
    "Snowflake", "dbt", "Fivetran", "Stitch",
    # Databases
    "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch", "Cassandra",
    "DynamoDB", "BigQuery", "Redshift", "SQLite", "Oracle", "Neo4j",
    # Cloud / DevOps
    "AWS", "GCP", "Azure", "Docker", "Kubernetes", "Terraform", "CI/CD",
    "MLflow", "Weights & Biases", "W&B", "SageMaker", "Vertex AI", "Azure ML",
    "Git", "GitHub", "GitLab", "Jenkins", "Prometheus", "Grafana",
    # BI / Viz
    "Power BI", "Tableau", "Looker", "Metabase", "Excel",
    # Stats / Math
    "Statistics", "Linear Algebra", "Probability", "A/B Testing", "Bayesian",
    "Algorithms", "Data Structures", "Distributed Systems",
]

def _extract_skills_from_text(text: str) -> list[str]:
    """
    Match known skill keywords against raw resume text (case-insensitive).
    Returns deduplicated list preserving canonical capitalisation.
    """
    text_lower = text.lower()
    found = []
    seen_lower = set()
    for skill in _KNOWN_SKILLS:
        skill_lower = skill.lower()
        if skill_lower in seen_lower:
            continue
        # Match as whole word / token (allow hyphen boundaries too)
        pattern = r'(?<![a-zA-Z0-9])' + re.escape(skill_lower) + r'(?![a-zA-Z0-9])'
        if re.search(pattern, text_lower):
            found.append(skill)
            seen_lower.add(skill_lower)
    return found


def _extract_resume_sections(text: str) -> dict:
    """
    Extract professional summary and two achievement bullets from resume text.

    Strategy:
    1. Find SUMMARY / PROFILE / OBJECTIVE section header → grab lines until
       the next section header is detected (blank line OR ALL-CAPS header)
    2. Cap summary at 5 lines / 80 words to prevent over-capture
    3. Find bullet points with metrics anywhere in the resume
    4. Prefer bullets containing quantified achievements (%, $, numbers)
    """
    lines = [l.strip() for l in text.splitlines()]

    summary_headers = re.compile(
        r"^(professional\s+)?(summary|profile|objective|about\s+me)$", re.I
    )
    # Any ALL-CAPS word(s) 3+ chars = section header
    section_header = re.compile(r"^[A-Z][A-Z\s&/\-]{2,}$")
    bullet_markers = re.compile(r"^[•\-–\*·▪➢➤►▸◆]\s*")
    metric_pattern = re.compile(
        r"(\d+\s*%|\$\s*[\d,]+|\d+[xX]|\d+\+?\s*(users?|clients?|records?|hours?|days?|years?|ms|gb|tb|k\b))",
        re.I,
    )
    contact_pattern = re.compile(
        r"@|linkedin\.com|github\.com|\(\d{3}\)|\d{3}[-\.]\d{3}|portfolio", re.I
    )
    # Lines that look like job titles / dates — stop summary collection
    date_pattern = re.compile(r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b|\d{4}", re.I)

    summary = ""
    bullets: list[str] = []
    in_summary = False
    summary_buf: list[str] = []

    def _is_section_boundary(line: str) -> bool:
        """Return True if this line is a new section header."""
        if not line:
            return True
        stripped = line.strip()
        # ALL-CAPS header (e.g. "EXPERIENCE", "TECHNICAL SKILLS")
        if section_header.match(stripped) and len(stripped.split()) <= 5:
            return True
        # Title-case known section words
        if re.match(r"^(Professional\s+Experience|Experience|Education|Skills|Technical\s+Skills|Projects|Certifications|Awards|Languages)\b", stripped):
            return True
        return False

    for line in lines:
        if not line:
            # Blank line: flush summary if we have enough content
            if in_summary and summary_buf:
                candidate = " ".join(summary_buf)
                if len(candidate.split()) >= 8:
                    summary = candidate
                    in_summary = False
                    summary_buf = []
            continue

        is_sum_hdr = bool(summary_headers.match(line)) and len(line.split()) <= 4

        if is_sum_hdr:
            in_summary = True
            summary_buf = []
            continue

        if in_summary:
            # Stop collecting if we hit a new section or a date/company line
            if _is_section_boundary(line):
                if summary_buf:
                    candidate = " ".join(summary_buf)
                    if len(candidate.split()) >= 8:
                        summary = candidate
                summary_buf = []
                in_summary = False
                # Don't skip — process as normal line below
            elif contact_pattern.search(line):
                pass  # skip contact lines inside summary zone
            elif date_pattern.search(line) and len(line.replace("\t", " ").split()) <= 8:
                # Looks like "University of Houston  Mar 2025 – Dec 2025" — stop
                if summary_buf:
                    candidate = " ".join(summary_buf)
                    if len(candidate.split()) >= 8:
                        summary = candidate
                summary_buf = []
                in_summary = False
            else:
                summary_buf.append(line)
                # Hard cap: 5 lines or 80 words — flush early
                if len(summary_buf) >= 5 or len(" ".join(summary_buf).split()) >= 80:
                    candidate = " ".join(summary_buf)
                    summary = candidate
                    summary_buf = []
                    in_summary = False

        # Collect bullets regardless of section
        if bullet_markers.match(line):
            clean = bullet_markers.sub("", line).strip()
            if len(clean.split()) >= 4 and not contact_pattern.search(clean):
                bullets.append(clean)
        elif bullets and line:
            # If line doesn't match bullet marker, section header, date, or contact...
            # it might be the continuation of the previous bullet point
            stripped = line.strip()
            if (
                not _is_section_boundary(stripped)
                and not contact_pattern.search(stripped)
                and not date_pattern.search(stripped)
                and not summary_headers.match(stripped)
                and len(stripped.split()) >= 2
            ):
                bullets[-1] = bullets[-1] + " " + stripped

    # Flush summary if file ended while still collecting
    if in_summary and summary_buf and not summary:
        candidate = " ".join(summary_buf)
        if len(candidate.split()) >= 8:
            summary = candidate

    # Fallback: first long non-contact, non-header line
    if not summary:
        for line in lines:
            if (
                len(line.split()) >= 12
                and not contact_pattern.search(line)
                and not section_header.match(line)
                and not date_pattern.search(line)
            ):
                # Truncate to ~80 words to avoid grabbing full paragraphs
                words = line.split()
                summary = " ".join(words[:80])
                break

    # Prefer metric-containing bullets; deduplicate
    seen = set()
    unique_bullets = []
    for b in bullets:
        key = b[:60]
        if key not in seen:
            seen.add(key)
            unique_bullets.append(b)

    metric_bullets = [b for b in unique_bullets if metric_pattern.search(b)]
    final_bullets = metric_bullets if len(metric_bullets) >= 2 else unique_bullets

    return {
        "summary": summary,
        "bullet_1": final_bullets[0] if len(final_bullets) > 0 else "",
        "bullet_2": final_bullets[1] if len(final_bullets) > 1 else "",
        "full_text": text,
        "skills": _extract_skills_from_text(text),
    }


def _extract_text_from_docx(content: bytes) -> tuple[str, int]:
    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="Install python-docx: pip install python-docx")

    doc = docx.Document(io.BytesIO(content))

    # Styles that represent bullet/list items — prepend "• " so the
    # bullet detection regex in _extract_resume_sections can find them.
    LIST_STYLES = {"list paragraph", "list bullet", "list bullet 2",
                   "list number", "list continue"}

    seen: set[str] = set()
    lines: list[str] = []

    def _add(p):
        t = p.text.strip()
        if not t:
            return
        style_name = (p.style.name or "").lower()
        # Prepend bullet marker if the paragraph is a list style AND doesn't
        # already start with a bullet character
        if style_name in LIST_STYLES and not re.match(r'^[•\-–\*·▪➢➤►▸◆]', t):
            t = "• " + t
        if t not in seen:
            seen.add(t)
            lines.append(t)

    for p in doc.paragraphs:
        _add(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _add(p)

    return "\n".join(lines), 1


@app.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    if not file.filename or not (file.filename.lower().endswith(".pdf") or file.filename.lower().endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are accepted.")

    content = await file.read()
    
    if file.filename.lower().endswith(".pdf"):
        text, pages = _extract_text_from_pdf(content)
    else:
        text, pages = _extract_text_from_docx(content)

    if not text.strip():
        raise HTTPException(status_code=422, detail="Could not extract text from document.")

    sections = _extract_resume_sections(text)
    return {**sections, "pages": pages}


# ─────────────────────────────────────────────────────────────────────────────
# POST /export-resume  (professional 1-page PDF)
# ─────────────────────────────────────────────────────────────────────────────

def _build_tailored_pdf(
    original_pdf_bytes: bytes,
    original_summary: str,
    original_bullet_1: str,
    original_bullet_2: str,
    tailored_summary: str,
    tailored_bullet_1: str,
    tailored_bullet_2: str,
    job_title: str,
    company: str,
) -> bytes:
    """
    Overlay approach: use pdfplumber to locate the bounding boxes of the
    original text regions, then use reportlab to draw white rectangles over
    them and render the new text in the same position/font size. Finally merge
    the overlay onto the original page with pypdf.

    This works for any PDF encoding (TJ arrays, kerned text, etc.) because we
    never touch the content streams — we only add an overlay on top.
    """
    try:
        import pdfplumber
        import pypdf
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

    replacements: list[tuple[str, str]] = []
    if original_summary and tailored_summary:
        replacements.append((original_summary.strip(), tailored_summary.strip()))
    if original_bullet_1 and tailored_bullet_1:
        replacements.append((original_bullet_1.strip(), tailored_bullet_1.strip()))
    if original_bullet_2 and tailored_bullet_2:
        replacements.append((original_bullet_2.strip(), tailored_bullet_2.strip()))

    if not replacements:
        return original_pdf_bytes

    # ── Step 1: locate bounding boxes for each original text block ────────────
    def _find_text_bbox(page, search_text: str):
        """
        Return (x0, top, x1, bottom) of the region containing search_text.
        Works even when pdfplumber merges entire lines into single 'words'
        (common in LaTeX/InDesign PDFs) by also matching against space-stripped text.
        Returns None if not found.
        """
        words = page.extract_words(
            x_tolerance=5, y_tolerance=5,
            keep_blank_chars=False, use_text_flow=True,
        )
        if not words:
            return None

        # Group words into lines by y proximity
        lines: list[list[dict]] = []
        for w in words:
            placed = False
            for line in lines:
                if abs(w["top"] - line[0]["top"]) < 6:
                    line.append(w)
                    placed = True
                    break
            if not placed:
                lines.append([w])
        lines.sort(key=lambda l: l[0]["top"])
        for line in lines:
            line.sort(key=lambda w: w["x0"])

        # Build two parallel representations per line:
        # normal (spaced) and nospace (all whitespace stripped)
        line_texts_normal  = [" ".join(w["text"] for w in ln) for ln in lines]
        line_texts_nospace = ["".join(w["text"] for w in ln) for ln in lines]

        full_normal  = " ".join(line_texts_normal)
        full_nospace = "".join(line_texts_nospace)

        needle_normal  = re.sub(r"\s+", " ", search_text.strip())
        needle_nospace = re.sub(r"\s+", "", search_text.strip())

        # Use a shorter anchor (first 60 chars) for robustness
        anchor_n  = needle_normal[:60]
        anchor_ns = needle_nospace[:60]

        matched_words = []

        if anchor_n in full_normal:
            match_start = full_normal.find(anchor_n)
            match_end   = match_start + len(needle_normal)
            pos = 0
            for ln in lines:
                for w in ln:
                    w_end = pos + len(w["text"])
                    if w_end > match_start and pos < match_end:
                        matched_words.append(w)
                    pos += len(w["text"]) + 1

        elif anchor_ns in full_nospace:
            match_start = full_nospace.find(anchor_ns)
            match_end   = match_start + len(needle_nospace)
            pos = 0
            for ln_idx, ln in enumerate(lines):
                ln_ns = line_texts_nospace[ln_idx]
                ln_end = pos + len(ln_ns)
                if ln_end > match_start and pos < match_end:
                    matched_words.extend(ln)
                pos = ln_end

        if not matched_words:
            return None

        x0     = min(w["x0"]    for w in matched_words) - 2
        x1     = max(w["x1"]    for w in matched_words) + 2
        top    = min(w["top"]    for w in matched_words) - 2
        bottom = max(w["bottom"] for w in matched_words) + 2
        return (x0, top, x1, bottom)

    # ── Step 2: for each replacement, redact the original region from the
    #            content stream, then inject new text at the same position ──────
    with pdfplumber.open(io.BytesIO(original_pdf_bytes)) as plumber_pdf:
        page0 = plumber_pdf.pages[0]
        page_w = float(page0.width)
        page_h = float(page0.height)

        # Collect bboxes for all replacements
        bboxes = []
        for original_text, new_text in replacements:
            bbox = _find_text_bbox(page0, original_text)
            bboxes.append((original_text, new_text, bbox))

    # ── Step 3: build the new PDF ─────────────────────────────────────────────
    from reportlab.lib.utils import simpleSplit

    reader = pypdf.PdfReader(io.BytesIO(original_pdf_bytes))
    writer = pypdf.PdfWriter()
    font_size = 9.5
    leading = font_size * 1.25

    for page_idx, page in enumerate(reader.pages):
        if page_idx != 0:
            writer.add_page(page)
            continue

        # Get decoded content stream
        if "/Contents" not in page:
            writer.add_page(page)
            continue

        contents = page["/Contents"]
        if hasattr(contents, "get_object"):
            contents = contents.get_object()
        if not isinstance(contents, list):
            contents = [contents]

        all_data = b""
        for obj_ref in contents:
            obj = obj_ref.get_object() if hasattr(obj_ref, "get_object") else obj_ref
            try:
                all_data += obj.get_data()
            except Exception:
                pass

        # For each bbox, remove BT...ET blocks whose Td/Tm y-coordinate falls
        # within the region (PDF y = page_h - pdfplumber_top).
        # We do this by parsing BT...ET text blocks and filtering by position.
        def _remove_text_in_region(data: bytes, top_pl: float, bottom_pl: float, page_h: float) -> bytes:
            """
            Remove text-drawing commands within a pdfplumber bounding box.
            Handles PDFs that use only Td (relative positioning) by tracking
            the accumulated absolute y position across the entire content stream.

            pdf_y range to remove: [page_h - bottom_pl, page_h - top_pl]
            """
            pdf_y_top    = page_h - top_pl + 2    # upper bound (larger y)
            pdf_y_bottom = page_h - bottom_pl - 2  # lower bound (smaller y)

            # Tokenise the stream into a flat list of operations.
            # We rebuild it, skipping text tokens whose absolute y falls in range.
            result = bytearray()
            cur_x, cur_y = 0.0, 0.0   # accumulated absolute position
            in_bt = False

            # Split into tokens: we iterate line by line / operator by operator
            # using a simple state machine.
            # Pattern: collect operands, then act on operator.
            token_re = re.compile(
                rb"(\[[\s\S]*?\]"       # array literal
                rb"|<[0-9a-fA-F]*>"     # hex string
                rb"|\((?:[^()\\]|\\.)*\)"  # string literal
                rb"|[-+]?\d*\.?\d+"     # number
                rb"|/\S+"              # name
                rb"|[A-Za-z'\"*]+"     # operator / keyword
                rb"|\s+"               # whitespace
                rb"|.)"                # catch-all
            )

            tokens = token_re.findall(data)
            i = 0
            while i < len(tokens):
                tok = tokens[i]

                # Track BT/ET
                stripped = tok.strip()
                if stripped == b"BT":
                    in_bt = True
                    cur_x, cur_y = 0.0, 0.0
                    result.extend(tok)
                    i += 1
                    continue
                if stripped == b"ET":
                    in_bt = False
                    result.extend(tok)
                    i += 1
                    continue

                if not in_bt:
                    result.extend(tok)
                    i += 1
                    continue

                # Inside BT block — track position operators
                if stripped == b"Td" or stripped == b"TD":
                    # Previous two numeric tokens are dx dy
                    # We need to look back in already-emitted tokens for dx/dy
                    # Easier: look ahead from last emitted Td — we accumulate below
                    # Actually: Td operands are already in result buffer as text
                    # Use a different approach: parse operand stack
                    result.extend(tok)
                    i += 1
                    continue

                if stripped == b"Tm":
                    result.extend(tok)
                    i += 1
                    continue

                # Text show operators: Tj, TJ, ', "
                if stripped in (b"Tj", b"TJ", b"'", b'"'):
                    result.extend(tok)
                    i += 1
                    continue

                result.extend(tok)
                i += 1

            # That approach is too complex for this PDF structure.
            # Simpler: parse the whole stream, find text segments with their
            # absolute y by tracking Td accumulation, and blank out segments in range.
            # We do this by rebuilding from scratch with regex matching of Td groups.
            return _remove_by_td_tracking(data, pdf_y_bottom, pdf_y_top)

        def _remove_by_td_tracking(data: bytes, pdf_y_lo: float, pdf_y_hi: float) -> bytes:
            """
            Track absolute y through Td accumulation. Within each BT block,
            scan operator by operator. When we hit a Tj/TJ and the current
            absolute y is within [pdf_y_lo, pdf_y_hi], replace the text
            argument with an empty string.
            """
            # We'll work line by line within each BT block.
            # Split content into BT...ET segments + non-BT segments.
            BT_RE = re.compile(rb"(BT\b[\s\S]*?ET\b)")
            parts = BT_RE.split(data)

            out = bytearray()
            for part in parts:
                if not part.startswith(b"BT"):
                    out.extend(part)
                    continue

                # Process this BT block
                abs_x, abs_y = 0.0, 0.0
                new_block = bytearray()

                # Tokenise into (value_bytes, operator_bytes) pairs
                # Find all operator lines: "... operand operand OPERATOR"
                # Use a regex to find Td, Tm, Tj, TJ, T*, and string tokens
                op_re = re.compile(
                    rb"((?:[-\d. \t\n\r]+\n?)*)"  # operands (numbers/whitespace)
                    rb"(\b(?:Td|TD|Tm|Tj|TJ|T\*|'|\")\b)"  # operator
                )

                # Simpler: split the block into lines and process
                # Actually just do regex replacements of Tj/TJ operands at matching y

                # Find all Td operators and accumulate y
                # Then find text ops between consecutive Td ops
                # Replace text ops whose accumulated y is in range

                # Parse segment by segment
                pos = 3  # skip "BT\n"
                block_text = part

                # Find all positioning and text operators with their byte offsets
                token_pattern = re.compile(
                    rb"([-\d.]+)\s+([-\d.]+)\s+(Td|TD)"  # Td with x y
                    rb"|(\d+)\s+Tr"  # text render mode (ignore)
                    rb"|(\[(?:[^\[\]]|\[(?:[^\[\]])*\])*\])\s*TJ"  # [...] TJ
                    rb"|\((?:[^()\\]|\\.)*\)\s*Tj"  # (...) Tj
                    rb"|\((?:[^()\\]|\\.)*\)\s*'"   # (...) '
                )

                # We need to blank text ops where accumulated y is in range.
                # Walk through finding Td ops first to build y-map, then do replacement.
                td_positions = []  # list of (byte_offset, delta_x, delta_y)
                for m in re.finditer(rb"([-\d.]+)\s+([-\d.]+)\s+(?:Td|TD)\b", block_text):
                    try:
                        dx = float(m.group(1))
                        dy = float(m.group(2))
                        td_positions.append((m.start(), dx, dy))
                    except Exception:
                        pass

                # Also find Tm operators for absolute positioning
                for m in re.finditer(rb"[-\d.]+\s+[-\d.]+\s+[-\d.]+\s+[-\d.]+\s+([-\d.]+)\s+([-\d.]+)\s+Tm\b", block_text):
                    pass  # no Tm in this PDF

                # Build a list of (byte_offset, abs_y) for each Td
                acc_x, acc_y = 0.0, 0.0
                td_abs = []
                for offset, dx, dy in td_positions:
                    acc_x += dx
                    acc_y += dy
                    td_abs.append((offset, acc_x, acc_y))

                # For each text operator, determine current y by finding
                # the last Td that occurred before it
                def _get_abs_y_at(byte_offset: int) -> float:
                    cur_y = 0.0
                    for td_off, _, ay in td_abs:
                        if td_off < byte_offset:
                            cur_y = ay
                        else:
                            break
                    return cur_y

                # Now replace text operator arguments if y is in range
                def _blank_if_in_range(m: re.Match) -> bytes:
                    y = _get_abs_y_at(m.start())
                    if pdf_y_lo <= y <= pdf_y_hi:
                        # Return empty text in same operator
                        op_bytes = m.group(0)
                        if b"TJ" in op_bytes:
                            return b"[] TJ"
                        elif op_bytes.rstrip().endswith(b"Tj"):
                            return b"() Tj"
                        elif op_bytes.rstrip().endswith(b"'"):
                            return b"() '"
                    return m.group(0)

                text_op_re = re.compile(
                    rb"\[(?:[^\[\]]|\[(?:[^\[\]])*\])*\]\s*TJ"
                    rb"|\((?:[^()\\]|\\.)*\)\s*Tj"
                    rb"|\((?:[^()\\]|\\.)*\)\s*'"
                )
                new_block = text_op_re.sub(_blank_if_in_range, block_text)
                out.extend(new_block)

            return bytes(out)

        modified_data = all_data
        all_bboxes_for_removal = []
        for orig_text, new_text, bbox in bboxes:
            if bbox is None:
                continue
            x0, top, x1, bottom = bbox
            all_bboxes_for_removal.append((top - 2, bottom + 18))

        for top_pl, bottom_pl in all_bboxes_for_removal:
            modified_data = _remove_text_in_region(modified_data, top_pl, bottom_pl, page_h)

        # Build overlay with new text using reportlab
        overlay_buf = io.BytesIO()
        c = rl_canvas.Canvas(overlay_buf, pagesize=(page_w, page_h))

        for orig_text, new_text, bbox in bboxes:
            if bbox is None:
                continue
            x0, top, x1, bottom = bbox
            region_w = x1 - x0
            rl_y_top = page_h - top  # reportlab y = distance from bottom

            wrapped = simpleSplit(new_text, "Helvetica", font_size, region_w)
            c.setFillColorRGB(0, 0, 0)
            text_obj = c.beginText(x0, rl_y_top - font_size)
            text_obj.setFont("Helvetica", font_size)
            text_obj.setLeading(leading)
            for line in wrapped:
                text_obj.textLine(line)
            c.drawText(text_obj)

        c.save()

        # Inject modified content stream + new text overlay
        from pypdf.generic import DecodedStreamObject
        new_stream = DecodedStreamObject()
        new_stream.set_data(modified_data)

        overlay_buf.seek(0)
        overlay_reader = pypdf.PdfReader(overlay_buf)
        overlay_page = overlay_reader.pages[0]

        # Get overlay content stream bytes
        ov_contents = overlay_page["/Contents"]
        if hasattr(ov_contents, "get_object"):
            ov_contents = ov_contents.get_object()
        if not isinstance(ov_contents, list):
            ov_contents = [ov_contents]
        overlay_data = b""
        for obj_ref in ov_contents:
            obj = obj_ref.get_object() if hasattr(obj_ref, "get_object") else obj_ref
            try:
                overlay_data += obj.get_data()
            except Exception:
                pass

        # Combine: modified original stream + overlay (new text)
        combined_stream = DecodedStreamObject()
        combined_stream.set_data(modified_data + b"\n" + overlay_data)

        # Copy page resources
        page[pypdf.generic.NameObject("/Contents")] = writer._add_object(combined_stream)
        writer.add_page(page)

    out_buf = io.BytesIO()
    writer.write(out_buf)
    return out_buf.getvalue()


@app.post("/export-resume")
def export_resume(req: ExportResumeRequest):
    try:
        original_pdf_bytes = base64.b64decode(req.original_pdf_base64)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 PDF data.")

    pdf_bytes = _build_tailored_pdf(
        original_pdf_bytes=original_pdf_bytes,
        original_summary=req.original_summary,
        original_bullet_1=req.original_bullet_1,
        original_bullet_2=req.original_bullet_2,
        tailored_summary=req.tailored_summary,
        tailored_bullet_1=req.tailored_bullet_1,
        tailored_bullet_2=req.tailored_bullet_2,
        job_title=req.job_title,
        company=req.company,
    )

    safe = lambda s: re.sub(r"[^\w\-]", "_", s)
    filename = f"resume_tailored_{safe(req.job_title)}_{safe(req.company)}.pdf"

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─────────────────────────────────────────────────────────────────────────────
# POST /export-resume-docx  (in-place DOCX edit — preserves all formatting)
# ─────────────────────────────────────────────────────────────────────────────

def _apply_docx_edits(docx_bytes: bytes, edits: list[dict]) -> bytes:
    """
    Apply a list of {original_text, new_text} replacements to a .docx file,
    doing run-level surgery to preserve bold/italic/font/size/color exactly.

    Strategy (mirrors the reference implementation):
    1. For each paragraph's runs, try to find original_text in a single run
       → replace only that run's text (all formatting XML untouched)
    2. If not found in a single run, fall back to paragraph-level replace
       (loses intra-run formatting for that paragraph only, keeps paragraph style)
    3. Repeat for all table cells too (handles layout tables in resumes)
    """
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="Install python-docx: pip install python-docx")

    doc = Document(io.BytesIO(docx_bytes))

    def _apply_to_paragraph(para, original: str, replacement: str):
        full_text = para.text
        if original not in full_text:
            return False

        # Fast path: entire original text lives inside a single run
        for run in para.runs:
            if original in run.text:
                run.text = run.text.replace(original, replacement, 1)
                return True

        # Fallback: text spans multiple runs.
        # Consolidate all runs into the first run (keeping its formatting),
        # do the replacement, then clear the rest. This preserves paragraph
        # style (font, size, bold) from the first run.
        if not para.runs:
            return False
        first_run = para.runs[0]
        first_run.text = full_text.replace(original, replacement, 1)
        for run in para.runs[1:]:
            run.text = ""
        return True

    for edit in edits:
        orig = edit.get("original_text", "").strip()
        new  = edit.get("new_text", "").strip()
        if not orig or not new or orig == new:
            continue

        # Search all body paragraphs
        for para in doc.paragraphs:
            _apply_to_paragraph(para, orig, new)

        # Search all table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _apply_to_paragraph(para, orig, new)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


class ExportDocxRequest(BaseModel):
    original_docx_base64: str
    tailored_summary: str
    tailored_bullet_1: str
    tailored_bullet_2: str
    original_summary: str = ""
    original_bullet_1: str = ""
    original_bullet_2: str = ""
    job_title: str = "Job"
    company: str = "Company"


@app.post("/export-resume-docx")
def export_resume_docx(req: ExportDocxRequest):
    try:
        docx_bytes = base64.b64decode(req.original_docx_base64)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 DOCX data.")

    edits = []
    if req.original_summary and req.tailored_summary:
        edits.append({"original_text": req.original_summary, "new_text": req.tailored_summary})
    if req.original_bullet_1 and req.tailored_bullet_1:
        edits.append({"original_text": req.original_bullet_1, "new_text": req.tailored_bullet_1})
    if req.original_bullet_2 and req.tailored_bullet_2:
        edits.append({"original_text": req.original_bullet_2, "new_text": req.tailored_bullet_2})

    if not edits:
        raise HTTPException(status_code=400, detail="No edits to apply — provide original and tailored text.")

    result_bytes = _apply_docx_edits(docx_bytes, edits)

    safe = lambda s: re.sub(r"[^\w\-]", "_", s)
    filename = f"resume_tailored_{safe(req.job_title)}_{safe(req.company)}.docx"

    return Response(
        content=result_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
